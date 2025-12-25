# /core/services/doc_reader_service.py
"""
Document Reader Service: อ่าน PDF และ DOC/DOCX
รองรับเอกสารขนาดใหญ่สำหรับ Bulk Import
"""

import io
from typing import Tuple, Optional
from pathlib import Path


class DocReaderService:
    """Service สำหรับอ่านเอกสาร PDF และ DOC/DOCX"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx'}
    
    def extract_text(self, file_bytes: bytes, filename: str) -> Tuple[str, int]:
        """
        อ่านข้อความจากไฟล์เอกสาร
        
        Args:
            file_bytes: เนื้อหาไฟล์ในรูป bytes
            filename: ชื่อไฟล์ (ใช้ระบุ extension)
            
        Returns:
            Tuple of (extracted_text, page_count)
        """
        ext = Path(filename).suffix.lower()
        
        if ext == '.pdf':
            return self._extract_from_pdf(file_bytes)
        elif ext in {'.doc', '.docx'}:
            return self._extract_from_docx(file_bytes)
        else:
            raise ValueError(f"ไม่รองรับประเภทไฟล์: {ext} รองรับเฉพาะ: {self.SUPPORTED_EXTENSIONS}")
    
    def _extract_from_pdf(self, file_bytes: bytes) -> Tuple[str, int]:
        """Extract text from PDF using PyMuPDF"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            page_count = len(doc)
            
            text_parts = []
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"--- หน้า {page_num} ---\n{page_text}")
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            print(f"✅ [DocReader] PDF: {page_count} หน้า, {len(full_text)} ตัวอักษร")
            return full_text, page_count
            
        except Exception as e:
            print(f"❌ [DocReader] ข้อผิดพลาด PDF: {e}")
            raise ValueError(f"ไม่สามารถอ่าน PDF ได้: {e}")
    
    def _extract_from_docx(self, file_bytes: bytes) -> Tuple[str, int]:
        """Extract text from DOC/DOCX using python-docx"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_bytes))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            # Estimate page count (rough: ~500 words per page)
            word_count = len(full_text.split())
            page_count = max(1, word_count // 500)
            
            print(f"✅ [DocReader] DOCX: ~{page_count} หน้า, {len(full_text)} ตัวอักษร")
            return full_text, page_count
            
        except Exception as e:
            print(f"❌ [DocReader] ข้อผิดพลาด DOCX: {e}")
            raise ValueError(f"ไม่สามารถอ่าน DOC/DOCX ได้: {e}")
    
    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS


# Singleton instance
doc_reader_service = DocReaderService()
